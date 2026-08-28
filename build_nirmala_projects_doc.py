from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "deliverables/Nirmala_Sundarappa_projects_and_activities.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
URL_PREFIX = "https://confluence.oraclecorp.com/confluence/pages/viewpage.action?pageId="

def set_font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), BLUE); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), "Calibri"); rFonts.set(qn("w:hAnsi"), "Calibri"); rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20"); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_link_line(doc, page_id, classification):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.1
    add_hyperlink(p, "Open in Oracle Central Confluence", URL_PREFIX + page_id)
    if classification:
        r = p.add_run(f"  |  {classification}")
        set_font(r, 9, MUTED, italic=True)

def add_priority(doc, num, title, text, page_id, classification):
    p = doc.add_paragraph(style="Priority Item")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{num}. {title}")
    set_font(r, 12, INK, bold=True)
    p2 = doc.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(1)
    add_link_line(doc, page_id, classification)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1); section.bottom_margin = Inches(1)
section.left_margin = Inches(1); section.right_margin = Inches(1)
section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    s = doc.styles[name]; s.font.name = "Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color); s.font.bold = True
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True

priority_style = doc.styles.add_style("Priority Item", WD_STYLE_TYPE.PARAGRAPH)
priority_style.paragraph_format.space_before = Pt(8); priority_style.paragraph_format.space_after = Pt(2)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("ORACLE CENTRAL CONFLUENCE  |  RESEARCH SUMMARY")
set_font(r, 8, MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Prioritized project and activity inventory")
set_font(r, 8, MUTED)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
r = title.add_run("Nirmala Sundarappa")
set_font(r, 25, INK, bold=True)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run("Projects and activities — prioritized for customer and external-developer impact")
set_font(r, 13, MUTED)

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(10)
r = lead.add_run("Scope. "); set_font(r, 11, INK, bold=True)
r = lead.add_run("This curated inventory is based on 50 Confluence search matches. It leads with direct customer engagements and external developer programs, then lists developer enablement and supporting internal references. Incidental name-only mentions are excluded.")
set_font(r, 11)

doc.add_heading("Priority: customer and external developer engagement", level=1)
priorities = [
    ("GDP for Java 2024", "Nirmala-led Global Developers Program for Java: selected customer and developer feedback, roadmap review, 23ai and open-source demos, polls, feature requests, and Q&A.", "10868324790", "Oracle Restricted"),
    ("Database Java/JDBC Customer Advisory Board", "Nirmala is the listed contact for the Database Java/JDBC Customer Advisory Board, providing a direct channel for structured customer input.", "2888149037", "Oracle Restricted"),
    ("FY26 customer engagements and developer evangelism", "The FY26 report records five customer engagements for Nirmala, including GDP4Java and Salesforce work, plus a developer-facing JDBC dependency-management blog.", "18841075375", "Oracle Internal"),
    ("Finanz Informatika + IBM Open Liberty proof of concept", "Customer proof of concept for UCP integration with IBM Open Liberty, with Nirmala on the development team and coordination involving IBM.", "2501857606", "Oracle Internal"),
    ("Workday Prism / Adaptive Planning HikariCP engagement", "Customer-facing technical work concerning HikariCP request boundaries and UCP migration recommendations for high-availability integration.", "11125798812", "Oracle Internal"),
    ("Oracle CloudWorld 2022", "Java-developer sessions and materials, including JDBC and Autonomous Database connectivity content.", "3587125186", None),
    ("Approved Sessions and LiveLabs", "CloudWorld sessions, JavaOne demos, and developer-facing event content with Nirmala listed as a speaker or contributor.", "4358375569", None),
]
for i, item in enumerate(priorities, 1): add_priority(doc, i, *item)

doc.add_heading("Developer education, adoption, and release outreach", level=1)
items = [
    ("Blogs (Published and Archived)", "Published Oracle Developers articles, including Spring Boot/UCP and JDBC/UCP release material.", "4434833845", "Oracle Internal"),
    ("Java/JDBC Blog & Technical Papers", "Planned technical content on Application Continuity, JDBC use cases, and Spring Boot with sharded databases.", "4181406268", "Oracle Internal"),
    ("DB 23ai Features Release Plan", "Code samples, blogs, and videos for JDBC/UCP capabilities.", "7517661524", "Oracle Internal"),
    ("Social media outreach", "JDBC tips and community posts for the Java Oracle Database audience.", "11139758233", "Oracle Internal"),
    ("Release Links and Release Checklist", "Developer-consumable release work: FAQ, Maven Central, OTN videos and technical briefs, and release messaging.", "2730559358", None),
    ("Developer-awareness profile", "Describes Nirmala’s role in customer feature adoption and her blogs, briefs, videos, and articles for the Java community.", "5939958479", "Oracle Internal"),
]
for title_text, activity, pid, cls in items:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title_text + ": "); set_font(r, 11, INK, bold=True)
    r = p.add_run(activity); set_font(r, 11)
    add_link_line(doc, pid, cls)

doc.add_heading("Supporting planning and references", level=1)
support = [
    ("O.com Migration — Website Builder", "AppDev and website-migration ownership reference.", "18176513324", "Oracle Restricted"),
    ("List of Oracle IDE Plugins", "Developer-tooling reference naming Nirmala and Kuassi.", "13035787410", None),
    ("2025 DBJava team notes", "Names Nirmala as the product-management contact for JDBC questions.", "13103301280", "Oracle Internal"),
    ("2 Day + Java Developer’s Guide planning", "Early Java documentation planning that includes customer-feedback and CAB work with Nirmala and Kuassi.", "241822042", None),
]
for title_text, activity, pid, cls in support:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title_text + ": "); set_font(r, 11, INK, bold=True)
    r = p.add_run(activity); set_font(r, 11)
    add_link_line(doc, pid, cls)

doc.save(OUT)
print(OUT)
